#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Converte il report da Markdown a Word (.docx).

Non e' un convertitore Markdown universale: gestisce solo quello che il nostro
report usa davvero (titoli, paragrafi, tabelle, elenchi, grassetto e corsivo).
E' una scelta voluta: meno codice, nessuna dipendenza esterna oltre python-docx,
e nessuna sorpresa su costrutti che non useremo mai.

Uso:
    python3 md_to_docx.py report_2026-07-15_2026-07-30.md
    python3 md_to_docx.py report.md nome_a_scelta.docx
"""

import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("ERRORE: manca la libreria python-docx.")
    print("Installala con:  pip install python-docx")
    sys.exit(1)


# ---------------------------------------------------------------- aspetto ---
FONT = "Calibri"
COLORE_TITOLO = RGBColor(0x1F, 0x36, 0x4D)      # blu scuro
COLORE_SEZIONE = RGBColor(0x2E, 0x74, 0xB5)     # blu medio
COLORE_INTESTAZIONE_TAB = "D9E2F3"              # azzurro chiaro (sfondo celle)
COLORE_ALLERTA = RGBColor(0xC0, 0x00, 0x00)     # rosso per le scadenze scadute

# Parole che fanno diventare rossa una riga di tabella.
SPIE_ALLERTA = ("in ritardo", "scaduta", "scaduto", "critica")


def sfondo_cella(cella, colore_hex):
    """Colora lo sfondo di una cella: python-docx non lo espone, serve l'XML."""
    tcPr = cella._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), colore_hex)
    tcPr.append(shd)


def scrivi_testo_formattato(paragrafo, testo, rosso=False):
    """
    Aggiunge testo a un paragrafo interpretando **grassetto** e *corsivo*.
    Divide la stringa sui marcatori e crea un "run" per ogni pezzo.
    """
    # Il gruppo di cattura tiene i separatori dentro il risultato dello split.
    pezzi = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)", testo)
    for pezzo in pezzi:
        if not pezzo:
            continue
        if pezzo.startswith("**") and pezzo.endswith("**"):
            run = paragrafo.add_run(pezzo[2:-2])
            run.bold = True
        elif pezzo.startswith("*") and pezzo.endswith("*"):
            run = paragrafo.add_run(pezzo[1:-1])
            run.italic = True
        elif pezzo.startswith("`") and pezzo.endswith("`"):
            run = paragrafo.add_run(pezzo[1:-1])
            run.font.name = "Consolas"
        else:
            run = paragrafo.add_run(pezzo)
        run.font.name = FONT
        if rosso:
            run.font.color.rgb = COLORE_ALLERTA


def righe_tabella(blocco):
    """Trasforma le righe markdown '| a | b |' in liste di celle."""
    righe = []
    for riga in blocco:
        # Scarta la riga separatrice |---|---|
        if re.fullmatch(r"\|[\s:|-]+\|", riga.strip()):
            continue
        celle = [c.strip() for c in riga.strip().strip("|").split("|")]
        righe.append(celle)
    return righe


def aggiungi_tabella(doc, blocco):
    dati = righe_tabella(blocco)
    if not dati:
        return
    n_colonne = max(len(r) for r in dati)
    tabella = doc.add_table(rows=0, cols=n_colonne)
    tabella.style = "Table Grid"
    tabella.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, riga_dati in enumerate(dati):
        celle = tabella.add_row().cells
        # Una riga e' "in allerta" se contiene una delle spie (es. scadenza scaduta)
        testo_riga = " ".join(riga_dati).lower()
        allerta = i > 0 and any(s in testo_riga for s in SPIE_ALLERTA)

        for j in range(n_colonne):
            contenuto = riga_dati[j] if j < len(riga_dati) else ""
            cella = celle[j]
            par = cella.paragraphs[0]
            # La cella nasce con un run vuoto: lo togliamo, altrimenti resta
            # un frammento di testo senza formattazione dentro il documento.
            for run_vuoto in list(par.runs):
                run_vuoto._element.getparent().remove(run_vuoto._element)
            par.paragraph_format.space_before = Pt(2)
            par.paragraph_format.space_after = Pt(2)
            if i == 0:
                # Intestazione: grassetto su sfondo azzurro
                run = par.add_run(re.sub(r"\*\*", "", contenuto))
                run.bold = True
                run.font.name = FONT
                run.font.size = Pt(10)
                sfondo_cella(cella, COLORE_INTESTAZIONE_TAB)
            else:
                scrivi_testo_formattato(par, contenuto, rosso=allerta)
                for run in par.runs:
                    run.font.size = Pt(10)


def converti(percorso_md, percorso_docx=None):
    percorso_md = Path(percorso_md)
    if not percorso_md.exists():
        print(f"ERRORE: non trovo il file {percorso_md}")
        return None
    if percorso_docx is None:
        percorso_docx = percorso_md.with_suffix(".docx")
    percorso_docx = Path(percorso_docx)

    testo = percorso_md.read_text(encoding="utf-8")
    # Via i commenti HTML (li usa solo _modello.md, ma meglio essere prudenti)
    testo = re.sub(r"<!--.*?-->", "", testo, flags=re.DOTALL)
    righe = testo.split("\n")

    doc = Document()
    stile = doc.styles["Normal"]
    stile.font.name = FONT
    stile.font.size = Pt(11)
    for sezione in doc.sections:
        sezione.left_margin = Cm(2.2)
        sezione.right_margin = Cm(2.2)

    i = 0
    while i < len(righe):
        riga = righe[i].rstrip()

        # --- riga vuota ---
        if not riga.strip():
            i += 1
            continue

        # --- tabella: raccolgo tutte le righe consecutive che iniziano con | ---
        if riga.lstrip().startswith("|"):
            blocco = []
            while i < len(righe) and righe[i].lstrip().startswith("|"):
                blocco.append(righe[i])
                i += 1
            aggiungi_tabella(doc, blocco)
            doc.add_paragraph()
            continue

        # --- titoli ---
        intestazione = re.match(r"^(#{1,4})\s+(.*)", riga)
        if intestazione:
            livello = len(intestazione.group(1))
            contenuto = intestazione.group(2).strip()
            # Usiamo i veri stili Titolo di Word (non solo grassetto piu' grande):
            # cosi' il documento ha una struttura navigabile e si puo' generare
            # un sommario automatico. Poi sovrascriviamo font, corpo e colore.
            par = doc.add_heading(level=livello)
            run = par.add_run(contenuto)
            run.font.name = FONT
            run.font.bold = True
            if livello == 1:
                run.font.size = Pt(20)
                run.font.color.rgb = COLORE_TITOLO
                par.paragraph_format.space_after = Pt(14)
            else:
                run.font.size = Pt(14 if livello == 2 else 12)
                run.font.color.rgb = COLORE_SEZIONE
                par.paragraph_format.space_before = Pt(12)
                par.paragraph_format.space_after = Pt(6)
            i += 1
            continue

        # --- elenco puntato ---
        punto = re.match(r"^[-*]\s+(.*)", riga)
        if punto:
            par = doc.add_paragraph(style="List Bullet")
            scrivi_testo_formattato(par, punto.group(1))
            i += 1
            continue

        # --- elenco numerato ---
        numero = re.match(r"^\d+\.\s+(.*)", riga)
        if numero:
            par = doc.add_paragraph(style="List Number")
            scrivi_testo_formattato(par, numero.group(1))
            i += 1
            continue

        # --- riga orizzontale ---
        if re.fullmatch(r"-{3,}|_{3,}|\*{3,}", riga.strip()):
            i += 1
            continue

        # --- paragrafo normale ---
        par = doc.add_paragraph()
        par.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        scrivi_testo_formattato(par, riga.strip())
        i += 1

    doc.save(percorso_docx)
    return percorso_docx


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print(__doc__)
        sys.exit(1)
    uscita = sys.argv[2] if len(sys.argv) > 2 else None
    risultato = converti(sys.argv[1], uscita)
    if risultato:
        print(risultato.name)
    else:
        sys.exit(1)
